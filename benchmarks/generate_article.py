# -*- coding: utf-8 -*-
"""Generate article .docx from content JSON."""

import json
import os

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

CONTENT_PATH = os.path.join(os.path.dirname(__file__), "article_content.json")
OUTPUT = os.path.join(os.path.dirname(__file__), "..", "..", "article_nn.docx")

FONT_NAME = "Times New Roman"
FONT_BODY = Pt(12)
FONT_TABLE = Pt(10)
FONT_LIT = Pt(11)


def set_font(run, size=FONT_BODY, bold=False, italic=False):
    run.font.name = FONT_NAME
    run.font.size = size
    run.font.bold = bold
    run.font.italic = italic
    rPr = run._element.find(qn("w:rPr"))
    if rPr is None:
        rPr = run._element.makeelement(qn("w:rPr"), {})
        run._element.insert(0, rPr)
    rF = rPr.find(qn("w:rFonts"))
    if rF is None:
        rF = rPr.makeelement(qn("w:rFonts"), {})
        rPr.append(rF)
    rF.set(qn("w:eastAsia"), FONT_NAME)
    rF.set(qn("w:cs"), FONT_NAME)


def set_pf(p, align=WD_ALIGN_PARAGRAPH.JUSTIFY, indent=Cm(1)):
    pf = p.paragraph_format
    pf.alignment = align
    pf.first_line_indent = indent
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing = 1.0


def add_p(doc, text, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
          bold=False, italic=False, size=FONT_BODY, indent=Cm(1)):
    p = doc.add_paragraph()
    set_pf(p, align=align, indent=indent)
    r = p.add_run(text)
    set_font(r, size=size, bold=bold, italic=italic)
    return p


def add_mixed(doc, parts, align=WD_ALIGN_PARAGRAPH.JUSTIFY, indent=Cm(1)):
    p = doc.add_paragraph()
    set_pf(p, align=align, indent=indent)
    for text, bold, italic, size in parts:
        r = p.add_run(text)
        set_font(r, size=size, bold=bold, italic=italic)
    return p


def fmt_cell(cell, size=FONT_TABLE):
    for p in cell.paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pf = p.paragraph_format
        pf.space_before = Pt(1)
        pf.space_after = Pt(1)
        pf.line_spacing = 1.0
        pf.first_line_indent = Pt(0)
        for r in p.runs:
            set_font(r, size=size)


def add_borders(table):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tblPr = table._tbl.tblPr
    if tblPr is None:
        tblPr = table._tbl.makeelement(qn("w:tblPr"), {})
    borders = tblPr.makeelement(qn("w:tblBorders"), {})
    for bn in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = borders.makeelement(qn("w:" + bn), {
            qn("w:val"): "single", qn("w:sz"): "4",
            qn("w:space"): "0", qn("w:color"): "000000",
        })
        borders.append(b)
    tblPr.append(borders)


def shade(row, color="D9D9D9"):
    for cell in row.cells:
        tcPr = cell._tc.get_or_add_tcPr()
        sh = tcPr.makeelement(qn("w:shd"), {
            qn("w:fill"): color, qn("w:val"): "clear",
        })
        tcPr.append(sh)


def main():
    with open(CONTENT_PATH, "r", encoding="utf-8") as f:
        c = json.load(f)

    doc = Document()

    for sec in doc.sections:
        sec.top_margin = Cm(2.5)
        sec.bottom_margin = Cm(2.5)
        sec.left_margin = Cm(2)
        sec.right_margin = Cm(2)
        sec.header_distance = Cm(0)
        sec.footer_distance = Cm(0)

    st = doc.styles["Normal"]
    st.font.name = FONT_NAME
    st.font.size = FONT_BODY
    st.paragraph_format.space_before = Pt(0)
    st.paragraph_format.space_after = Pt(0)
    st.paragraph_format.line_spacing = 1.0

    L = WD_ALIGN_PARAGRAPH.LEFT
    R = WD_ALIGN_PARAGRAPH.RIGHT
    C = WD_ALIGN_PARAGRAPH.CENTER
    J = WD_ALIGN_PARAGRAPH.JUSTIFY
    Z = Pt(0)

    # UDK
    add_p(doc, c["udk"], align=L, indent=Z)

    # Authors
    add_p(doc, "", indent=Z)
    add_p(doc, c["authors"], align=R, indent=Z)
    add_p(doc, c["univ"], align=R, indent=Z, italic=True)

    # Title
    add_p(doc, "", indent=Z)
    add_p(doc, c["title"], align=C, bold=True, indent=Z)
    add_p(doc, "", indent=Z)

    # Introduction
    add_p(doc, c["intro1"])
    add_p(doc, c["intro2"])

    # Architecture
    add_mixed(doc, [
        (c["arch_title"], True, False, FONT_BODY),
        (c["arch_body"], False, False, FONT_BODY),
    ])

    add_mixed(doc, [
        (c["stream_title"], True, False, FONT_BODY),
        (c["stream_body"], False, False, FONT_BODY),
    ])

    add_mixed(doc, [
        (c["xref_title"], True, False, FONT_BODY),
        (c["xref_body"], False, False, FONT_BODY),
    ])

    # Table 1
    add_p(doc, "", indent=Z)
    add_p(doc, c["table1_title"], align=C, indent=Z, size=FONT_LIT)

    t1_data = c["table1"]
    t1 = doc.add_table(rows=len(t1_data), cols=3)
    add_borders(t1)
    for ri, row in enumerate(t1_data):
        for ci, val in enumerate(row):
            cell = t1.rows[ri].cells[ci]
            cell.text = val
            fmt_cell(cell)
            if ci == 0:
                for p in cell.paragraphs:
                    p.alignment = L
    shade(t1.rows[0])

    # Experimental section
    add_p(doc, "", indent=Z)
    add_mixed(doc, [
        (c["exp_title"], True, False, FONT_BODY),
        (c["exp_body"], False, False, FONT_BODY),
    ])

    # Table 2
    add_p(doc, "", indent=Z)
    add_p(doc, c["table2_title"], align=C, indent=Z, size=FONT_LIT)

    headers = c["table2_headers"]
    data = c["table2_data"]
    t2 = doc.add_table(rows=1 + len(data), cols=len(headers))
    add_borders(t2)

    for ci, h in enumerate(headers):
        cell = t2.rows[0].cells[ci]
        cell.text = h
        fmt_cell(cell)
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.bold = True
    shade(t2.rows[0])

    for ri, row in enumerate(data):
        for ci, val in enumerate(row):
            cell = t2.rows[ri + 1].cells[ci]
            cell.text = val
            fmt_cell(cell)
            if ci == 0:
                for p in cell.paragraphs:
                    p.alignment = L

    # Analysis
    add_p(doc, "", indent=Z)
    add_p(doc, c["analysis"])

    # Conclusions
    add_mixed(doc, [
        (c["conclusions_label"], False, True, FONT_BODY),
        (c["conclusions"], False, False, FONT_BODY),
    ])

    # References
    add_p(doc, "", indent=Z)
    add_p(doc, c["lit_label"], align=L, indent=Z, size=FONT_LIT)

    for ref in c["refs"]:
        add_p(doc, ref, align=J, indent=Z, size=FONT_LIT)

    doc.save(OUTPUT)
    print(f"Article saved to {OUTPUT}")
    print(f"Size: {os.path.getsize(OUTPUT) / 1024:.1f} KB")


if __name__ == "__main__":
    main()
